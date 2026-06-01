"""
Unit tests for DocxTranslationAdapter.

Tests the adapter implementation for the generic orchestrator.
"""

import os
import pytest
from unittest.mock import Mock, MagicMock

from src.core.docx.docx_translation_adapter import DocxTranslationAdapter


class TestDocxTranslationAdapter:
    """Tests for DocxTranslationAdapter class."""

    def test_init(self):
        """Test adapter initialization."""
        adapter = DocxTranslationAdapter()

        assert adapter is not None
        assert adapter.converter is not None
        assert adapter.container is not None
        assert adapter.tag_preserver is not None
        assert adapter.html_chunker is not None

    def test_extract_content(self, simple_docx_path):
        """Test content extraction from DOCX."""
        adapter = DocxTranslationAdapter()
        log_callback = Mock()

        html_content, context = adapter.extract_content(simple_docx_path, log_callback)

        # Check HTML content
        assert html_content is not None
        assert len(html_content) > 0
        assert 'paragraph' in html_content.lower()

        # Check context structure
        assert 'metadata' in context
        assert 'preserver' in context
        assert 'source_path' in context
        assert context['source_path'] == simple_docx_path

        # Check log callback was called
        log_callback.assert_called()

    def test_extract_content_no_callback(self, simple_docx_path):
        """Test extraction without log callback."""
        adapter = DocxTranslationAdapter()

        html_content, context = adapter.extract_content(simple_docx_path, None)

        # Should work without callback
        assert html_content is not None
        assert context is not None

    def test_preserve_structure(self, simple_docx_path):
        """Test structure preservation via placeholders."""
        adapter = DocxTranslationAdapter()
        log_callback = Mock()

        # First extract content
        html_content, context = adapter.extract_content(simple_docx_path, None)

        # Preserve structure
        text_with_placeholders, tag_map, placeholder_format = adapter.preserve_structure(
            html_content, context, log_callback
        )

        # Check placeholders
        assert text_with_placeholders is not None
        assert isinstance(tag_map, dict)
        assert isinstance(placeholder_format, tuple)
        assert len(placeholder_format) == 2

        # Check log callback
        log_callback.assert_called()

    def test_preserve_structure_formats_placeholders(self, formatted_docx_path):
        """Test placeholder format for tags."""
        adapter = DocxTranslationAdapter()

        html_content, context = adapter.extract_content(formatted_docx_path, None)
        text_with_placeholders, tag_map, placeholder_format = adapter.preserve_structure(
            html_content, context, None
        )

        # Should have placeholders for HTML tags
        prefix, suffix = placeholder_format
        assert prefix is not None
        assert suffix is not None

        # Tag map should have entries
        if tag_map:
            # Check format
            for placeholder_id, tag_content in tag_map.items():
                assert isinstance(placeholder_id, str)
                assert isinstance(tag_content, str)

    def test_create_chunks(self, simple_docx_path):
        """Test chunk creation."""
        adapter = DocxTranslationAdapter()
        log_callback = Mock()

        # Extract and preserve
        html_content, context = adapter.extract_content(simple_docx_path, None)
        text_with_placeholders, tag_map, _ = adapter.preserve_structure(html_content, context, None)

        # Create chunks
        chunks = adapter.create_chunks(text_with_placeholders, tag_map, 450, log_callback)

        # Check chunks
        assert isinstance(chunks, list)
        assert len(chunks) > 0

        # Check log callback
        log_callback.assert_called()

    def test_create_chunks_structure(self, formatted_docx_path):
        """Test chunk structure."""
        adapter = DocxTranslationAdapter()

        html_content, context = adapter.extract_content(formatted_docx_path, None)
        text_with_placeholders, tag_map, _ = adapter.preserve_structure(html_content, context, None)
        chunks = adapter.create_chunks(text_with_placeholders, tag_map, 450, None)

        # Each chunk should be a dict with expected fields
        for chunk in chunks:
            assert isinstance(chunk, dict)
            # HtmlChunker should return dicts with content
            assert 'content' in chunk or 'text' in chunk

    def test_reconstruct_content(self, simple_docx_path):
        """Test content reconstruction from chunks."""
        adapter = DocxTranslationAdapter()

        # Extract and preserve
        html_content, context = adapter.extract_content(simple_docx_path, None)
        text_with_placeholders, tag_map, _ = adapter.preserve_structure(html_content, context, None)

        # Create fake translated chunks
        chunks = adapter.create_chunks(text_with_placeholders, tag_map, 450, None)
        translated_chunks = [f"[TRANSLATED] {chunk.get('content', chunk.get('text', ''))}"
                           for chunk in chunks]

        # Reconstruct
        reconstructed = adapter.reconstruct_content(translated_chunks, tag_map, context)

        # Check reconstruction
        assert reconstructed is not None
        assert isinstance(reconstructed, str)
        assert len(reconstructed) > 0

    def test_reconstruct_restores_tags(self, formatted_docx_path):
        """Test that tags are restored correctly."""
        adapter = DocxTranslationAdapter()

        html_content, context = adapter.extract_content(formatted_docx_path, None)
        text_with_placeholders, tag_map, _ = adapter.preserve_structure(html_content, context, None)

        # Simulate translation (keep placeholders)
        translated_chunks = [text_with_placeholders]

        # Reconstruct
        reconstructed = adapter.reconstruct_content(translated_chunks, tag_map, context)

        # Should have HTML tags restored (if there were any)
        if '<' in html_content:
            # Tags should be restored
            assert isinstance(reconstructed, str)

    def test_finalize_output(self, simple_docx_path, temp_dir):
        """Test output finalization (HTML → DOCX bytes)."""
        adapter = DocxTranslationAdapter()
        log_callback = Mock()

        # Extract content
        html_content, context = adapter.extract_content(simple_docx_path, None)

        # Finalize (convert back to DOCX)
        docx_bytes = adapter.finalize_output(html_content, simple_docx_path, context, log_callback)

        # Check result
        assert docx_bytes is not None
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

        # Check log callback
        log_callback.assert_called()

        # Test that bytes are valid DOCX
        from docx import Document
        import io
        doc = Document(io.BytesIO(docx_bytes))
        assert doc is not None

    def test_finalize_output_creates_valid_docx(self, formatted_docx_path, temp_dir):
        """Test finalized DOCX is valid and contains content."""
        adapter = DocxTranslationAdapter()

        html_content, context = adapter.extract_content(formatted_docx_path, None)
        docx_bytes = adapter.finalize_output(html_content, formatted_docx_path, context, None)

        # Load as DOCX
        from docx import Document
        import io
        doc = Document(io.BytesIO(docx_bytes))

        # Check has content
        assert len(doc.paragraphs) > 0

    def test_full_pipeline(self, simple_docx_path):
        """Test full adapter pipeline."""
        adapter = DocxTranslationAdapter()

        # 1. Extract
        html_content, context = adapter.extract_content(simple_docx_path, None)
        assert html_content is not None

        # 2. Preserve
        text_with_placeholders, tag_map, placeholder_format = adapter.preserve_structure(
            html_content, context, None
        )
        assert text_with_placeholders is not None

        # 3. Chunk
        chunks = adapter.create_chunks(text_with_placeholders, tag_map, 450, None)
        assert len(chunks) > 0

        # 4. Simulate translation
        translated_chunks = [chunk.get('content', chunk.get('text', '')) for chunk in chunks]

        # 5. Reconstruct
        reconstructed = adapter.reconstruct_content(translated_chunks, tag_map, context)
        assert reconstructed is not None

        # 6. Finalize
        docx_bytes = adapter.finalize_output(reconstructed, simple_docx_path, context, None)
        assert docx_bytes is not None
        assert len(docx_bytes) > 0

    def test_pipeline_with_formatting(self, formatted_docx_path):
        """Test pipeline preserves formatting through full cycle."""
        adapter = DocxTranslationAdapter()

        # Extract
        html_content, context = adapter.extract_content(formatted_docx_path, None)

        # Check HTML has formatting tags
        original_has_tags = '<' in html_content

        # Preserve
        text_with_placeholders, tag_map, _ = adapter.preserve_structure(html_content, context, None)

        # Chunk
        chunks = adapter.create_chunks(text_with_placeholders, tag_map, 450, None)

        # Simulate translation (preserve structure)
        translated_chunks = [chunk.get('content', chunk.get('text', '')) for chunk in chunks]

        # Reconstruct
        reconstructed = adapter.reconstruct_content(translated_chunks, tag_map, context)

        # Finalize
        docx_bytes = adapter.finalize_output(reconstructed, formatted_docx_path, context, None)

        # Result should be valid DOCX
        assert isinstance(docx_bytes, bytes)
        assert len(docx_bytes) > 0

    def test_empty_content(self, temp_dir):
        """Test handling of empty content."""
        adapter = DocxTranslationAdapter()

        # Create empty DOCX
        from docx import Document
        doc = Document()
        empty_path = os.path.join(temp_dir, 'empty.docx')
        doc.save(empty_path)

        # Extract (should handle empty content)
        html_content, context = adapter.extract_content(empty_path, None)

        # May be empty or minimal HTML
        assert context is not None

    def test_error_handling_invalid_path(self):
        """Test error handling for invalid file path."""
        adapter = DocxTranslationAdapter()

        with pytest.raises(Exception):
            adapter.extract_content('/nonexistent/path.docx', None)

    def test_context_preservation(self, simple_docx_path):
        """Test that context is preserved through pipeline."""
        adapter = DocxTranslationAdapter()

        # Extract
        html_content, context = adapter.extract_content(simple_docx_path, None)

        # Context should have required fields
        assert 'metadata' in context
        assert 'preserver' in context
        assert 'source_path' in context

        # Context should be usable in preserve_structure
        text_with_placeholders, tag_map, _ = adapter.preserve_structure(html_content, context, None)

        # Context should be usable in reconstruct_content
        reconstructed = adapter.reconstruct_content([text_with_placeholders], tag_map, context)
        assert reconstructed is not None

        # Context should be usable in finalize_output
        docx_bytes = adapter.finalize_output(reconstructed, simple_docx_path, context, None)
        assert docx_bytes is not None
