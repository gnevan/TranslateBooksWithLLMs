"""
Pytest fixtures for DOCX translation tests.
"""

import os
import pytest
import tempfile
from docx import Document
from docx.shared import Pt, Inches


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir


@pytest.fixture
def simple_docx_path(temp_dir):
    """Create a simple DOCX file for testing."""
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a simple paragraph.')
    doc.add_paragraph('This is another paragraph with bold text.').runs[0].bold = True

    path = os.path.join(temp_dir, 'simple.docx')
    doc.save(path)
    return path


@pytest.fixture
def formatted_docx_path(temp_dir):
    """Create a DOCX with various formatting."""
    doc = Document()

    # Headings
    doc.add_heading('Main Title', 0)
    doc.add_heading('Section 1', 1)

    # Paragraph with mixed formatting
    p = doc.add_paragraph('This paragraph has ')
    p.add_run('bold').bold = True
    p.add_run(' and ')
    p.add_run('italic').italic = True
    p.add_run(' and ')
    p.add_run('underlined').underline = True
    p.add_run(' text.')

    # List
    doc.add_paragraph('First item', style='List Bullet')
    doc.add_paragraph('Second item', style='List Bullet')

    # Numbered list
    doc.add_paragraph('Item 1', style='List Number')
    doc.add_paragraph('Item 2', style='List Number')

    path = os.path.join(temp_dir, 'formatted.docx')
    doc.save(path)
    return path


@pytest.fixture
def table_docx_path(temp_dir):
    """Create a DOCX with a table."""
    doc = Document()
    doc.add_heading('Document with Table', 0)

    # Add table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    # Fill table
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f'Cell {i},{j}'

    path = os.path.join(temp_dir, 'with_table.docx')
    doc.save(path)
    return path


@pytest.fixture
def sample_html():
    """Sample HTML content for testing."""
    return """<html><body>
<h1>Test Title</h1>
<p>This is a simple paragraph.</p>
<p>This has <strong>bold</strong> and <em>italic</em> text.</p>
<ul>
<li>First item</li>
<li>Second item</li>
</ul>
</body></html>"""


@pytest.fixture
def sample_metadata():
    """Sample metadata for testing."""
    return {
        'styles': {},
        'default_font': {'name': 'Calibri', 'size': 11},
        'page_size': {'width': 8.5, 'height': 11.0},
        'margins': {
            'top': 1.0,
            'bottom': 1.0,
            'left': 1.0,
            'right': 1.0
        }
    }


@pytest.fixture
def mock_llm_client():
    """Mock LLM client for testing."""
    class MockLLMClient:
        async def translate(self, text, source_lang, target_lang, **kwargs):
            # Simple mock translation: add "[FR]" prefix
            return f"[FR] {text}"

    return MockLLMClient()
