"""
Unit tests for DocxHtmlConverter.

Tests DOCX ↔ HTML conversion functionality.
"""

import os
import pytest
from docx import Document
from lxml import etree

from src.core.docx.converter import DocxHtmlConverter


class TestDocxHtmlConverter:
    """Tests for DocxHtmlConverter class."""

    def test_init(self):
        """Test converter initialization."""
        converter = DocxHtmlConverter()
        assert converter is not None

    def test_to_html_simple(self, simple_docx_path):
        """Test conversion of simple DOCX to HTML."""
        converter = DocxHtmlConverter()
        html_content, metadata = converter.to_html(simple_docx_path)

        # Check HTML content is returned
        assert html_content is not None
        assert len(html_content) > 0
        assert 'simple paragraph' in html_content.lower()

        # Check metadata structure
        assert 'page_size' in metadata
        assert 'margins' in metadata
        assert 'default_font' in metadata

    def test_to_html_formatted(self, formatted_docx_path):
        """Test conversion preserves formatting."""
        converter = DocxHtmlConverter()
        html_content, metadata = converter.to_html(formatted_docx_path)

        # Check for headings
        assert '<h1>' in html_content or 'Main Title' in html_content

        # Check for formatted text (mammoth should produce semantic HTML)
        assert 'bold' in html_content.lower()
        assert 'italic' in html_content.lower()

    def test_to_html_with_table(self, table_docx_path):
        """Test conversion of DOCX with table."""
        converter = DocxHtmlConverter()
        html_content, metadata = converter.to_html(table_docx_path)

        # Mammoth may convert tables to HTML tables
        # Just check basic content is present
        assert 'Cell' in html_content
        assert html_content is not None

    def test_extract_metadata(self, simple_docx_path):
        """Test metadata extraction."""
        converter = DocxHtmlConverter()
        _, metadata = converter.to_html(simple_docx_path)

        # Check page size
        assert metadata['page_size'] is not None
        assert 'width' in metadata['page_size']
        assert 'height' in metadata['page_size']

        # Check margins
        assert metadata['margins'] is not None
        assert 'top' in metadata['margins']
        assert 'bottom' in metadata['margins']
        assert 'left' in metadata['margins']
        assert 'right' in metadata['margins']

        # Check default font
        assert metadata['default_font'] is not None
        assert 'name' in metadata['default_font']
        assert 'size' in metadata['default_font']

    def test_from_html_simple(self, temp_dir, sample_html, sample_metadata):
        """Test conversion from HTML to DOCX."""
        converter = DocxHtmlConverter()
        output_path = os.path.join(temp_dir, 'output.docx')

        # Convert HTML to DOCX
        converter.from_html(sample_html, sample_metadata, output_path)

        # Check file was created
        assert os.path.exists(output_path)

        # Load and check content
        doc = Document(output_path)
        text_content = '\n'.join([p.text for p in doc.paragraphs])

        assert 'Test Title' in text_content
        assert 'simple paragraph' in text_content

    def test_from_html_preserves_formatting(self, temp_dir, sample_metadata):
        """Test HTML to DOCX preserves basic formatting."""
        html = """<html><body>
<p>Normal text with <strong>bold</strong> and <em>italic</em> parts.</p>
</body></html>"""

        converter = DocxHtmlConverter()
        output_path = os.path.join(temp_dir, 'formatted_output.docx')

        converter.from_html(html, sample_metadata, output_path)

        # Load and check
        assert os.path.exists(output_path)
        doc = Document(output_path)

        # Check paragraph exists
        assert len(doc.paragraphs) > 0

        # Check runs exist (formatting splits text into runs)
        paragraph = doc.paragraphs[0]
        assert len(paragraph.runs) > 0

    def test_from_html_with_headings(self, temp_dir, sample_metadata):
        """Test HTML headings are converted to DOCX headings."""
        html = """<html><body>
<h1>Heading 1</h1>
<h2>Heading 2</h2>
<p>Regular paragraph.</p>
</body></html>"""

        converter = DocxHtmlConverter()
        output_path = os.path.join(temp_dir, 'headings_output.docx')

        converter.from_html(html, sample_metadata, output_path)

        assert os.path.exists(output_path)
        doc = Document(output_path)

        # Check paragraphs exist
        assert len(doc.paragraphs) >= 3

        # Check heading text
        text_content = '\n'.join([p.text for p in doc.paragraphs])
        assert 'Heading 1' in text_content
        assert 'Heading 2' in text_content

    def test_from_html_with_lists(self, temp_dir, sample_metadata):
        """Test HTML lists are converted to DOCX lists."""
        html = """<html><body>
<ul>
<li>First item</li>
<li>Second item</li>
</ul>
<ol>
<li>Numbered 1</li>
<li>Numbered 2</li>
</ol>
</body></html>"""

        converter = DocxHtmlConverter()
        output_path = os.path.join(temp_dir, 'lists_output.docx')

        converter.from_html(html, sample_metadata, output_path)

        assert os.path.exists(output_path)
        doc = Document(output_path)

        # Check list items exist
        text_content = '\n'.join([p.text for p in doc.paragraphs])
        assert 'First item' in text_content
        assert 'Second item' in text_content
        assert 'Numbered 1' in text_content

    def test_from_html_with_table(self, temp_dir, sample_metadata):
        """Test HTML table is converted to DOCX table."""
        html = """<html><body>
<table>
<tr><td>A1</td><td>A2</td></tr>
<tr><td>B1</td><td>B2</td></tr>
</table>
</body></html>"""

        converter = DocxHtmlConverter()
        output_path = os.path.join(temp_dir, 'table_output.docx')

        converter.from_html(html, sample_metadata, output_path)

        assert os.path.exists(output_path)
        doc = Document(output_path)

        # Check table exists
        assert len(doc.tables) > 0

        # Check table dimensions
        table = doc.tables[0]
        assert len(table.rows) == 2
        assert len(table.columns) == 2

        # Check cell content
        assert table.cell(0, 0).text == 'A1'
        assert table.cell(0, 1).text == 'A2'
        assert table.cell(1, 0).text == 'B1'
        assert table.cell(1, 1).text == 'B2'

    def test_apply_page_metadata(self, temp_dir, sample_metadata):
        """Test page metadata is applied correctly."""
        converter = DocxHtmlConverter()
        output_path = os.path.join(temp_dir, 'metadata_output.docx')

        html = "<html><body><p>Test</p></body></html>"
        converter.from_html(html, sample_metadata, output_path)

        # Load and check
        doc = Document(output_path)
        section = doc.sections[0]

        # Check page size (approximate comparison due to unit conversion)
        assert abs(section.page_width.inches - sample_metadata['page_size']['width']) < 0.1
        assert abs(section.page_height.inches - sample_metadata['page_size']['height']) < 0.1

        # Check margins
        assert abs(section.top_margin.inches - sample_metadata['margins']['top']) < 0.1
        assert abs(section.bottom_margin.inches - sample_metadata['margins']['bottom']) < 0.1

    def test_roundtrip_conversion(self, simple_docx_path, temp_dir):
        """Test DOCX → HTML → DOCX roundtrip."""
        converter = DocxHtmlConverter()

        # Convert to HTML
        html_content, metadata = converter.to_html(simple_docx_path)

        # Convert back to DOCX
        output_path = os.path.join(temp_dir, 'roundtrip.docx')
        converter.from_html(html_content, metadata, output_path)

        # Load original and converted
        original_doc = Document(simple_docx_path)
        converted_doc = Document(output_path)

        # Compare text content
        original_text = '\n'.join([p.text for p in original_doc.paragraphs])
        converted_text = '\n'.join([p.text for p in converted_doc.paragraphs])

        # Text should be preserved (formatting may differ slightly)
        assert 'simple paragraph' in converted_text.lower()

    def test_get_text_content(self):
        """Test _get_text_content helper method."""
        converter = DocxHtmlConverter()

        html = "<p>Text with <strong>bold</strong> part</p>"
        tree = etree.HTML(html)
        p_element = tree.find('.//p')

        text = converter._get_text_content(p_element)
        assert text == "Text with bold part"

    def test_empty_html(self, temp_dir, sample_metadata):
        """Test handling of empty HTML."""
        converter = DocxHtmlConverter()
        output_path = os.path.join(temp_dir, 'empty_output.docx')

        html = "<html><body></body></html>"
        converter.from_html(html, sample_metadata, output_path)

        # Should create valid DOCX even if empty
        assert os.path.exists(output_path)
        doc = Document(output_path)
        assert doc is not None

    def test_malformed_html(self, temp_dir, sample_metadata):
        """Test handling of malformed HTML."""
        converter = DocxHtmlConverter()
        output_path = os.path.join(temp_dir, 'malformed_output.docx')

        # HTML without proper structure
        html = "<p>Just a paragraph"
        converter.from_html(html, sample_metadata, output_path)

        # lxml should handle it gracefully
        assert os.path.exists(output_path)
