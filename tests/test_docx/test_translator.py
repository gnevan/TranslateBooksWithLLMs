"""
Integration tests for DOCX translation.

Tests the complete translate_docx_file function with mocked LLM.
"""

import os
import pytest
from unittest.mock import Mock, AsyncMock, MagicMock
from docx import Document

from src.core.docx.translator import translate_docx_file


class MockLLMProvider:
    """Mock LLM provider for testing."""

    def __init__(self, translation_prefix="[TRANSLATED]"):
        self.translation_prefix = translation_prefix
        self.call_count = 0
        self.model = "mock-model"

    async def generate(self, prompt, timeout=None, system_prompt=None, **kwargs):
        """Mock translation - returns LLMResponse."""
        from src.core.llm.base import LLMResponse
        self.call_count += 1
        # Extract text from prompt (simplified) - return text between tags
        content = f"{self.translation_prefix} Mock translation {self.call_count}"
        return LLMResponse(
            content=f"<TRANSLATION_START>{content}<TRANSLATION_END>",
            prompt_tokens=50,
            completion_tokens=20,
            context_used=70,
            context_limit=8000
        )

    def extract_translation(self, response):
        """Extract translation from response."""
        # Simple extraction - remove tags
        if "<TRANSLATION_START>" in response and "<TRANSLATION_END>" in response:
            start = response.index("<TRANSLATION_START>") + len("<TRANSLATION_START>")
            end = response.index("<TRANSLATION_END>")
            return response[start:end]
        return response

    async def send_prompt(self, prompt, **kwargs):
        """Mock translation - legacy method."""
        self.call_count += 1
        return f"{self.translation_prefix} Mock translation {self.call_count}"


@pytest.fixture
def mock_llm_provider():
    """Create mock LLM provider."""
    return MockLLMProvider()


class TestTranslateDocxFile:
    """Integration tests for translate_docx_file function."""

    @pytest.mark.asyncio
    async def test_translate_simple_docx(self, simple_docx_path, temp_dir, mock_llm_provider):
        """Test translation of simple DOCX file."""
        output_path = os.path.join(temp_dir, 'translated_simple.docx')

        result = await translate_docx_file(
            input_filepath=simple_docx_path,
            output_filepath=output_path,
            source_language='English',
            target_language='French',
            model_name='test-model',
            llm_client=mock_llm_provider,
            max_tokens_per_chunk=450,
            log_callback=None,
            progress_callback=None,
            prompt_options=None,
            max_retries=1
        )

        # Check result structure
        assert result['success'] is True
        assert 'stats' in result
        assert result['output_path'] == output_path

        # Check file was created
        assert os.path.exists(output_path)

        # Check file is valid DOCX
        doc = Document(output_path)
        assert doc is not None
        assert len(doc.paragraphs) > 0

    @pytest.mark.asyncio
    async def test_translate_formatted_docx(self, formatted_docx_path, temp_dir, mock_llm_provider):
        """Test translation preserves formatting."""
        output_path = os.path.join(temp_dir, 'translated_formatted.docx')

        result = await translate_docx_file(
            input_filepath=formatted_docx_path,
            output_filepath=output_path,
            source_language='English',
            target_language='Spanish',
            model_name='test-model',
            llm_client=mock_llm_provider,
            max_tokens_per_chunk=450
        )

        assert result['success'] is True
        assert os.path.exists(output_path)

        # Load and verify
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0

    @pytest.mark.asyncio
    async def test_translate_with_table(self, table_docx_path, temp_dir, mock_llm_provider):
        """Test translation of DOCX with table."""
        output_path = os.path.join(temp_dir, 'translated_table.docx')

        result = await translate_docx_file(
            input_filepath=table_docx_path,
            output_filepath=output_path,
            source_language='English',
            target_language='German',
            model_name='test-model',
            llm_client=mock_llm_provider
        )

        assert result['success'] is True
        assert os.path.exists(output_path)

        # Check table is preserved
        doc = Document(output_path)
        assert len(doc.tables) > 0

    @pytest.mark.asyncio
    async def test_translate_with_callbacks(self, simple_docx_path, temp_dir, mock_llm_provider):
        """Test translation with callbacks."""
        output_path = os.path.join(temp_dir, 'translated_callbacks.docx')

        log_callback = Mock()
        stats_callback = Mock()

        result = await translate_docx_file(
            input_filepath=simple_docx_path,
            output_filepath=output_path,
            source_language='English',
            target_language='French',
            model_name='test-model',
            llm_client=mock_llm_provider,
            log_callback=log_callback,
            stats_callback=stats_callback
        )

        assert result['success'] is True

        # Callbacks should have been called
        assert log_callback.call_count > 0
        assert stats_callback.call_count > 0

    @pytest.mark.asyncio
    async def test_translate_stats_tracking(self, simple_docx_path, temp_dir, mock_llm_provider):
        """Test that statistics are tracked correctly."""
        output_path = os.path.join(temp_dir, 'translated_stats.docx')

        result = await translate_docx_file(
            input_filepath=simple_docx_path,
            output_filepath=output_path,
            source_language='English',
            target_language='French',
            model_name='test-model',
            llm_client=mock_llm_provider
        )

        # Check stats structure
        assert 'stats' in result
        stats = result['stats']

        # Stats should have standard fields (from TranslationMetrics)
        assert isinstance(stats, dict)

    @pytest.mark.asyncio
    async def test_translate_with_refinement(self, simple_docx_path, temp_dir, mock_llm_provider):
        """Test translation with refinement option."""
        output_path = os.path.join(temp_dir, 'translated_refined.docx')

        prompt_options = {'refine': True}

        result = await translate_docx_file(
            input_filepath=simple_docx_path,
            output_filepath=output_path,
            source_language='English',
            target_language='French',
            model_name='test-model',
            llm_client=mock_llm_provider,
            prompt_options=prompt_options
        )

        assert result['success'] is True
        assert os.path.exists(output_path)

    @pytest.mark.asyncio
    async def test_translate_max_retries(self, simple_docx_path, temp_dir, mock_llm_provider):
        """Test max_retries parameter is passed through."""
        output_path = os.path.join(temp_dir, 'translated_retries.docx')

        result = await translate_docx_file(
            input_filepath=simple_docx_path,
            output_filepath=output_path,
            source_language='English',
            target_language='French',
            model_name='test-model',
            llm_client=mock_llm_provider,
            max_retries=3
        )

        assert result['success'] is True

    @pytest.mark.asyncio
    async def test_translate_custom_chunk_size(self, simple_docx_path, temp_dir, mock_llm_provider):
        """Test custom max_tokens_per_chunk."""
        output_path = os.path.join(temp_dir, 'translated_chunks.docx')

        result = await translate_docx_file(
            input_filepath=simple_docx_path,
            output_filepath=output_path,
            source_language='English',
            target_language='French',
            model_name='test-model',
            llm_client=mock_llm_provider,
            max_tokens_per_chunk=200  # Smaller chunks
        )

        assert result['success'] is True
        assert os.path.exists(output_path)

    @pytest.mark.asyncio
    async def test_translate_output_bytes_valid(self, simple_docx_path, temp_dir, mock_llm_provider):
        """Test that output file is valid DOCX."""
        output_path = os.path.join(temp_dir, 'translated_valid.docx')

        result = await translate_docx_file(
            input_filepath=simple_docx_path,
            output_filepath=output_path,
            source_language='English',
            target_language='French',
            model_name='test-model',
            llm_client=mock_llm_provider
        )

        # Check file exists and is readable
        assert os.path.exists(output_path)
        assert os.path.getsize(output_path) > 0

        # Verify it's a valid DOCX by loading it
        doc = Document(output_path)
        assert doc is not None

    @pytest.mark.asyncio
    async def test_translate_error_handling_invalid_input(self, temp_dir, mock_llm_provider):
        """Test error handling for invalid input file."""
        output_path = os.path.join(temp_dir, 'output.docx')

        with pytest.raises(Exception):
            await translate_docx_file(
                input_filepath='/nonexistent/input.docx',
                output_filepath=output_path,
                source_language='English',
                target_language='French',
                model_name='test-model',
                llm_client=mock_llm_provider
            )

    @pytest.mark.asyncio
    async def test_translate_preserves_page_settings(self, simple_docx_path, temp_dir, mock_llm_provider):
        """Test that page settings are preserved."""
        output_path = os.path.join(temp_dir, 'translated_page_settings.docx')

        # Get original page settings
        original_doc = Document(simple_docx_path)
        original_width = original_doc.sections[0].page_width
        original_height = original_doc.sections[0].page_height

        # Translate
        result = await translate_docx_file(
            input_filepath=simple_docx_path,
            output_filepath=output_path,
            source_language='English',
            target_language='French',
            model_name='test-model',
            llm_client=mock_llm_provider
        )

        assert result['success'] is True

        # Check page settings preserved (approximately)
        translated_doc = Document(output_path)
        translated_width = translated_doc.sections[0].page_width
        translated_height = translated_doc.sections[0].page_height

        # Should be similar (allowing for small differences)
        assert abs(original_width - translated_width) < 100000  # In EMUs
        assert abs(original_height - translated_height) < 100000

    @pytest.mark.asyncio
    async def test_translate_with_context_manager(self, simple_docx_path, temp_dir, mock_llm_provider):
        """Test translation with context manager."""
        output_path = os.path.join(temp_dir, 'translated_context.docx')

        # Mock context manager
        context_manager = Mock()

        result = await translate_docx_file(
            input_filepath=simple_docx_path,
            output_filepath=output_path,
            source_language='English',
            target_language='French',
            model_name='test-model',
            llm_client=mock_llm_provider,
            context_manager=context_manager
        )

        assert result['success'] is True

    @pytest.mark.asyncio
    async def test_translate_empty_docx(self, temp_dir, mock_llm_provider):
        """Test translation of empty DOCX."""
        # Create empty DOCX
        empty_doc = Document()
        empty_path = os.path.join(temp_dir, 'empty.docx')
        empty_doc.save(empty_path)

        output_path = os.path.join(temp_dir, 'translated_empty.docx')

        result = await translate_docx_file(
            input_filepath=empty_path,
            output_filepath=output_path,
            source_language='English',
            target_language='French',
            model_name='test-model',
            llm_client=mock_llm_provider
        )

        # Should succeed even with empty content
        assert result['success'] is True
        assert os.path.exists(output_path)

    @pytest.mark.asyncio
    async def test_translate_multiple_paragraphs(self, temp_dir, mock_llm_provider):
        """Test translation of DOCX with multiple paragraphs."""
        # Create DOCX with many paragraphs
        doc = Document()
        for i in range(10):
            doc.add_paragraph(f'Paragraph {i + 1} with some content to translate.')

        input_path = os.path.join(temp_dir, 'multi_para.docx')
        doc.save(input_path)

        output_path = os.path.join(temp_dir, 'translated_multi_para.docx')

        result = await translate_docx_file(
            input_filepath=input_path,
            output_filepath=output_path,
            source_language='English',
            target_language='French',
            model_name='test-model',
            llm_client=mock_llm_provider
        )

        assert result['success'] is True
        assert os.path.exists(output_path)

        # Check all paragraphs present
        translated_doc = Document(output_path)
        assert len(translated_doc.paragraphs) >= 10

    @pytest.mark.asyncio
    async def test_orchestrator_integration(self, simple_docx_path, temp_dir, mock_llm_provider):
        """Test that orchestrator is used correctly."""
        output_path = os.path.join(temp_dir, 'translated_orchestrator.docx')

        # The function should use GenericTranslationOrchestrator internally
        result = await translate_docx_file(
            input_filepath=simple_docx_path,
            output_filepath=output_path,
            source_language='English',
            target_language='French',
            model_name='test-model',
            llm_client=mock_llm_provider
        )

        # Should complete successfully via orchestrator
        assert result['success'] is True
        assert os.path.exists(output_path)
        assert 'stats' in result
        assert 'output_path' in result
